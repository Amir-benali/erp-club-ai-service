from __future__ import annotations

from functools import lru_cache
from io import BytesIO
from typing import Any, Dict, List, Optional, Tuple
import os
import re

import numpy as np
from PIL import Image, ImageEnhance, ImageOps

try:
    from rapidocr_onnxruntime import RapidOCR
except Exception as exc:  # pragma: no cover
    RapidOCR = None
    RAPIDOCR_IMPORT_ERROR = str(exc)
else:
    RAPIDOCR_IMPORT_ERROR = None

try:
    import easyocr
except Exception as exc:  # pragma: no cover
    easyocr = None
    EASYOCR_IMPORT_ERROR = str(exc)
else:
    EASYOCR_IMPORT_ERROR = None

try:
    import pytesseract
except Exception as exc:  # pragma: no cover
    pytesseract = None
    TESSERACT_IMPORT_ERROR = str(exc)
else:
    TESSERACT_IMPORT_ERROR = None


NUTRIENT_ALIASES: Dict[str, List[str]] = {
    "vitamin_d": ["vitamin d", "vit d", "vitamine d", "vit. d", "25-oh d", "25 oh d", "25(oh)d", "25-oh", "25 oh"],
    "vitamin_b12": ["vitamin b12", "vit b12", "vitamine b12", "vit. b12", "cobalamin", "cobalamine"],
    "folate": ["folate", "folic acid", "acide folique", "vitamin b9", "vitamine b9", "vit b9"],
    "vitamin_c": ["vitamin c", "vitamine c", "vit c", "ascorbic acid", "acide ascorbique"],
    "vitamin_a": ["vitamin a", "vitamine a", "retinol", "rétinol"],
    "vitamin_e": ["vitamin e", "vitamine e", "tocopherol", "tocophérol"],
    "zinc": ["zinc"],
    "magnesium": ["magnesium", "magnésium", "magnesuim", "magnisium", "magnesiumm", "magnesium serum", "magnésium sérique"],
    "iron": ["iron", "fer", "fe", "serum iron", "fer sérique"],
    "ferritin": ["ferritin", "ferritine"],
    "calcium": ["calcium", "ca"],
}

# Common markers in Tunisian French sports-medical reports.  The endpoint keeps
# the legacy ``nutrients_found`` name, but includes these clinical biomarkers.
BIOMARKER_ALIASES: Dict[str, List[str]] = {
    "hemoglobin": ["hémoglobine", "hemoglobine", "hgb", "hb"],
    "c_reactive_protein": ["protéine c-réactive", "proteine c-reactive", "crp", "crp ultrasensible"],
    "fasting_glucose": ["glycémie à jeun", "glycemie a jeun", "glycémie", "glycemie", "glucose à jeun", "glucose fasting"],
    "total_cholesterol": ["cholestérol total", "cholesterol total"],
    "ldl_cholesterol": ["ldl-cholestérol", "ldl-cholesterol", "ldl cholesterol"],
    "hdl_cholesterol": ["hdl-cholestérol", "hdl-cholesterol", "hdl cholesterol"],
    "triglycerides": ["triglycérides", "triglycerides"],
}

REFERENCE_RANGES: Dict[str, Dict[str, Tuple[float, float]]] = {
    "vitamin_d": {"ng/ml": (30.0, 100.0)},
    "vitamin_b12": {"pg/ml": (200.0, 900.0)},
    "folate": {"ng/ml": (3.0, 17.0)},
    "zinc": {"ug/dl": (70.0, 120.0)},
    "magnesium": {"mg/dl": (1.7, 2.4)},
    "ferritin": {"ng/ml": (30.0, 400.0)},
    "calcium": {"mg/dl": (8.6, 10.2)},
    "hemoglobin": {"g/dl": (13.0, 17.0)},
    "c_reactive_protein": {"mg/l": (0.0, 3.0)},
    "fasting_glucose": {"g/l": (0.70, 1.10)},
    "total_cholesterol": {"g/l": (0.0, 2.00)},
    "ldl_cholesterol": {"g/l": (0.0, 1.50)},
    "hdl_cholesterol": {"g/l": (0.40, 10.0)},
    "triglycerides": {"g/l": (0.0, 1.50)},
}

VALUE_PATTERN = r"([<>]?\s*\d+(?:[\.,]\d+)?)"
UNIT_PATTERN = r"(ng\s*/\s*ml|ng\s*/\s*l|pg\s*/\s*ml|mg\s*/\s*dl|mg\s*/\s*l|g\s*/\s*dl|g\s*/\s*l|mmol\s*/\s*l|[uµμ]mol\s*/\s*l|[uµμ]g\s*/\s*dl|mcg\s*/\s*l)?"
REQUIRED_UNIT_PATTERN = UNIT_PATTERN[:-1]

# Some PDF generators emit a table by column rather than by row. In the report
# format used by Tunisian clinics, these results may therefore appear directly
# before their label in the extracted text.
VALUE_BEFORE_LABEL_MARKERS = {
    "total_cholesterol",
    "ldl_cholesterol",
    "triglycerides",
    "vitamin_b12",
    "magnesium",
    "c_reactive_protein",
}


@lru_cache(maxsize=1)
def get_rapidocr_reader() -> Any:
    if RapidOCR is None:
        return None
    return RapidOCR()


@lru_cache(maxsize=1)
def get_easyocr_reader() -> Any:
    if easyocr is None:
        return None
    return easyocr.Reader(["fr", "en"], gpu=False)


def run_easyocr_on_image_bytes(image_bytes: bytes) -> str:
    image = Image.open(BytesIO(image_bytes)).convert("RGB")
    # Medical reports are often photographed; enlarging and increasing contrast
    # improves French accents, decimal commas, and small table values.
    image = ImageOps.grayscale(image)
    image = ImageEnhance.Contrast(image).enhance(1.5)
    # Avoid turning high-resolution scans into images too large for the OCR
    # detector. Small phone images are enlarged; large PDF renders are capped.
    # Cloud containers can be memory-constrained. A 2,800 px edge provides
    # readable medical tables without allocating hundreds of MB in ONNX/OpenCV.
    max_dimension = int(os.getenv("OCR_MAX_DIMENSION", "2800"))
    scale = min(2.0, max_dimension / max(image.width, image.height))
    if scale != 1.0:
        image = image.resize(
            (max(1, int(image.width * scale)), max(1, int(image.height * scale))),
            Image.Resampling.LANCZOS,
        )

    # Tesseract has a far smaller memory footprint than ONNX Runtime. It is
    # selected by the production Docker image for constrained Railway plans.
    backend = os.getenv("OCR_BACKEND", "rapidocr").strip().lower()
    if backend == "tesseract" and pytesseract is not None:
        try:
            text = pytesseract.image_to_string(
                image,
                lang="fra+eng",
                config="--oem 1 --psm 11",
            )
            # Do not initialize ONNX as a fallback in a constrained cloud
            # container: it is precisely the allocation that can terminate
            # the service. An empty result becomes a normal API 400 response.
            return text.strip()
        except Exception as exc:
            raise RuntimeError(f"Tesseract OCR failed: {exc}") from exc

    np_image = np.array(image)

    rapid_reader = get_rapidocr_reader()
    if rapid_reader is not None:
        result, _ = rapid_reader(np_image)
        if result:
            lines = [item[1] for item in result if len(item) >= 2]
            text = "\n".join([str(line).strip() for line in lines if str(line).strip()])
            if text.strip():
                return text

    easy_reader = get_easyocr_reader()
    if easy_reader is not None:
        lines = easy_reader.readtext(np_image, detail=0, paragraph=True)
        return "\n".join([str(line).strip() for line in lines if str(line).strip()])

    raise RuntimeError(
        "No OCR backend is available. "
        f"RapidOCR import error: {RAPIDOCR_IMPORT_ERROR or 'not installed'}; "
        f"EasyOCR import error: {EASYOCR_IMPORT_ERROR or 'not installed'}; "
        f"Tesseract import error: {TESSERACT_IMPORT_ERROR or 'not installed'}."
    )


# ── PDF / DOCX text extraction ──────────────────────────────────────────

def _merge_unique_lines(*text_sources: str) -> str:
    """Merge PDF text sources without repeating the same visible line."""
    seen = set()
    merged: List[str] = []
    for source in text_sources:
        for line in (source or "").splitlines():
            cleaned = " ".join(line.split())
            key = re.sub(r"[^\w]+", "", cleaned.lower())
            if cleaned and key and key not in seen:
                seen.add(key)
                merged.append(cleaned)
    return "\n".join(merged)


def _render_pdf_page_for_ocr(page: Any, file_bytes: bytes, page_index: int) -> bytes:
    """Render one PDF page to PNG with pdfplumber, then pypdfium2 as fallback."""
    try:
        preview = page.to_image(resolution=220)
        image_buffer = BytesIO()
        preview.original.save(image_buffer, format="PNG")
        return image_buffer.getvalue()
    except Exception:
        # pypdfium2 handles many scanned/image-only PDFs that pdfplumber's
        # renderer cannot decode or rasterize.
        import pypdfium2 as pdfium

        pdf_document = pdfium.PdfDocument(file_bytes)
        pdf_page = pdf_document[page_index]
        bitmap = pdf_page.render(scale=220 / 72)
        image_buffer = BytesIO()
        bitmap.to_pil().save(image_buffer, format="PNG")
        return image_buffer.getvalue()


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract embedded text and OCR every rendered PDF page.

    Clinic systems frequently generate PDFs with a partial text layer: headings
    are selectable, but table values are drawn as glyphs or images.  Combining
    native extraction with page OCR prevents those values from being dropped.
    """
    import pdfplumber

    text_parts: List[str] = []
    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        for page_index, page in enumerate(pdf.pages):
            native_text = page.extract_text(layout=True) or page.extract_text() or ""
            ocr_text = ""
            try:
                ocr_text = run_easyocr_on_image_bytes(
                    _render_pdf_page_for_ocr(page, file_bytes, page_index)
                )
            except Exception:
                # Native extraction is still useful when the renderer/OCR fails.
                pass

            page_text = _merge_unique_lines(native_text, ocr_text)
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    from docx import Document

    doc = Document(BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Route to the correct parser based on file extension."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext in ("pdf",):
        return extract_text_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(file_bytes)
    elif ext in ("png", "jpg", "jpeg", "tiff", "bmp", "webp"):
        return run_easyocr_on_image_bytes(file_bytes)
    else:
        raise ValueError(
            f"Unsupported file type: .{ext}. "
            "Supported: PDF, DOCX, PNG, JPG, JPEG, TIFF, BMP, WEBP."
        )


# ── Nutrient extraction ─────────────────────────────────────────────────
def _normalize_unit(unit: Optional[str]) -> Optional[str]:
    if not unit:
        return None
    return (
        unit.strip()
        .lower()
        .replace(" ", "")
        .replace("µ", "u")
        .replace("μ", "u")
    )


def _to_float(value_text: str) -> Optional[float]:
    if not value_text:
        return None
    cleaned = value_text.replace("<", "").replace(">", "").replace(" ", "")
    cleaned = cleaned.replace(",", ".")
    try:
        return float(cleaned)
    except ValueError:
        return None


def _status_from_reference(
    nutrient: str,
    value: Optional[float],
    unit: Optional[str],
) -> str:
    if value is None or not unit:
        return "unknown"

    nutrient_ranges = REFERENCE_RANGES.get(nutrient)
    if not nutrient_ranges:
        return "unknown"

    unit_range = nutrient_ranges.get(unit)
    if not unit_range:
        return "unknown"

    low, high = unit_range
    if value < low:
        return "low"
    if value > high:
        return "high"
    return "normal"


def extract_nutrients_from_text(text: str) -> Dict[str, Any]:
    lowered = text.lower()
    mentions: List[Dict[str, Any]] = []

    markers = {**NUTRIENT_ALIASES, **BIOMARKER_ALIASES}
    for nutrient, aliases in markers.items():
        for alias in aliases:
            alias_regex = re.escape(alias)
            if nutrient in VALUE_BEFORE_LABEL_MARKERS:
                # Select the closest unit-bearing value before the label. The
                # tempered section prevents an earlier result in another row
                # from being paired with this label.
                preceding_value_pattern = re.compile(
                    rf"{VALUE_PATTERN}\s*{REQUIRED_UNIT_PATTERN}"
                    rf"(?:(?!{VALUE_PATTERN}\s*{REQUIRED_UNIT_PATTERN})[\s\S]){{0,100}}?"
                    rf"\b{alias_regex}\b",
                    re.IGNORECASE,
                )
                preceding_matches = list(preceding_value_pattern.finditer(lowered))
                if preceding_matches:
                    match = preceding_matches[-1]
                    raw_value = match.group(1)
                    raw_unit = match.group(2)
                    value = _to_float(raw_value)
                    unit = _normalize_unit(raw_unit)
                    start, end = match.span()
                    context_start = max(0, start - 35)
                    context_end = min(len(text), end + 35)
                    mentions.append(
                        {
                            "nutrient": nutrient,
                            "matched_alias": alias,
                            "value": value,
                            "unit": unit,
                            "status": _status_from_reference(nutrient, value, unit),
                            "text_snippet": text[context_start:context_end].strip(),
                            "match_start": start,
                            "match_end": end,
                        }
                    )
                    continue

            # PDF table extraction often puts a reference interval immediately
            # after the parameter label (e.g. "Ferritine 30 - 400 85 ng/mL").
            # First collect a value that carries a unit; this is the measured
            # result, while the unitted numbers are normally reference bounds.
            unit_value_pattern = re.compile(
                rf"\b{alias_regex}\b[\s\S]{{0,140}}?{VALUE_PATTERN}\s*{REQUIRED_UNIT_PATTERN}",
                re.IGNORECASE,
            )
            for match in unit_value_pattern.finditer(lowered):
                raw_value = match.group(1)
                raw_unit = match.group(2)
                value = _to_float(raw_value)
                unit = _normalize_unit(raw_unit)
                start, end = match.span()
                context_start = max(0, start - 35)
                context_end = min(len(text), end + 35)
                mentions.append(
                    {
                        "nutrient": nutrient,
                        "matched_alias": alias,
                        "value": value,
                        "unit": unit,
                        "status": _status_from_reference(nutrient, value, unit),
                        "text_snippet": text[context_start:context_end].strip(),
                        "match_start": start,
                        "match_end": end,
                    }
                )

            pattern = re.compile(
                rf"\b{alias_regex}\b[^\n\r\d]{{0,40}}{VALUE_PATTERN}\s*{UNIT_PATTERN}",
                re.IGNORECASE,
            )

            for match in pattern.finditer(lowered):
                raw_value = match.group(1)
                raw_unit = match.group(2)
                value = _to_float(raw_value)
                unit = _normalize_unit(raw_unit)
                status = _status_from_reference(nutrient, value, unit)

                start, end = match.span()
                context_start = max(0, start - 35)
                context_end = min(len(text), end + 35)

                mentions.append(
                    {
                        "nutrient": nutrient,
                        "matched_alias": alias,
                        "value": value,
                        "unit": unit,
                        "status": status,
                        "text_snippet": text[context_start:context_end].strip(),
                        "match_start": start,
                        "match_end": end,
                    }
                )

    # Deduplicate by overlap + quality, then by nutrient/value/unit.
    def quality_score(item: Dict[str, Any]) -> int:
        score = 0
        if item.get("unit"):
            score += 3
        if item.get("value") is not None:
            score += 2
        if item.get("status") in {"low", "high", "normal"}:
            score += 2
        return score

    mentions_sorted = sorted(
        mentions,
        key=lambda x: (x["nutrient"], x["match_start"], -(x["match_end"] - x["match_start"])),
    )

    overlap_filtered: List[Dict[str, Any]] = []
    for item in mentions_sorted:
        replaced = False
        for idx, existing in enumerate(overlap_filtered):
            if item["nutrient"] != existing["nutrient"]:
                continue
            overlaps = not (
                item["match_end"] <= existing["match_start"]
                or item["match_start"] >= existing["match_end"]
            )
            if overlaps:
                if quality_score(item) > quality_score(existing):
                    overlap_filtered[idx] = item
                replaced = True
                break
        if not replaced:
            overlap_filtered.append(item)

    seen = set()
    deduped_mentions: List[Dict[str, Any]] = []
    for item in overlap_filtered:
        key = (item["nutrient"], item["value"], item["unit"])
        if key in seen:
            continue
        seen.add(key)
        deduped_mentions.append(item)

    for item in deduped_mentions:
        item.pop("match_start", None)
        item.pop("match_end", None)

    # If a nutrient has at least one unit-bearing mention, discard weaker
    # unit-less unknown mentions for that same nutrient.
    nutrients_with_unit = {
        m["nutrient"] for m in deduped_mentions if m.get("unit")
    }
    deduped_mentions = [
        m
        for m in deduped_mentions
        if not (
            m["nutrient"] in nutrients_with_unit
            and not m.get("unit")
            and m.get("status") == "unknown"
        )
    ]

    flagged = [m for m in deduped_mentions if m["status"] in {"low", "high"}]

    return {
        "mentions": deduped_mentions,
        "flagged": flagged,
        "nutrients_found": sorted({m["nutrient"] for m in deduped_mentions}),
    }
